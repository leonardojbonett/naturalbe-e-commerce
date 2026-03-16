#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Análisis detallado del DOCX: mejoras_seo_ids_855_881.docx
"""

try:
    from docx import Document
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False
    print("❌ python-docx no instalado. Instalando...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    DOC_AVAILABLE = True

from pathlib import Path

docx_path = Path('mejoras_seo_ids_855_881.docx')

print("🔍 Analizando estructura del DOCX...\n")

doc = Document(docx_path)

# Contar elementos
print(f"📊 Elementos en el documento:")
print(f"   Párrafos: {len(doc.paragraphs)}")
print(f"   Tablas: {len(doc.tables)}")
print()

# Analizar párrafos
print("📝 Primeros 10 párrafos:")
print("-" * 70)
for i, para in enumerate(doc.paragraphs[:10]):
    text = para.text.strip()
    if text:
        print(f"{i}: {text[:80]}")
print()

# Analizar tablas
if doc.tables:
    print(f"📋 Analizando {len(doc.tables)} tabla(s):")
    print("-" * 70)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTabla {table_idx + 1}:")
        print(f"  Filas: {len(table.rows)}")
        print(f"  Columnas: {len(table.columns)}")
        
        # Mostrar encabezados
        if table.rows:
            print(f"\n  Encabezados:")
            header_cells = table.rows[0].cells
            for col_idx, cell in enumerate(header_cells):
                print(f"    Col {col_idx}: {cell.text.strip()[:50]}")
        
        # Mostrar primeras 5 filas de datos
        print(f"\n  Primeras 5 filas de datos:")
        for row_idx, row in enumerate(table.rows[1:6]):
            print(f"\n  Fila {row_idx + 1}:")
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()[:80]
                print(f"    Col {col_idx}: {cell_text}")
else:
    print("❌ No hay tablas en el documento")
    print("\n📝 Analizando contenido de párrafos para datos...")
    
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"\nPárrafo {i}: {para.text[:100]}")
