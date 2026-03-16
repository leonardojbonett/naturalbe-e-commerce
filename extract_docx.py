#!/usr/bin/env python3
import json
from docx import Document

doc = Document('mejoras_seo_ids_824_854.docx')

# Write all paragraph text to JSON arrays for easy reading
data = {
    'paragraphs': [p.text for p in doc.paragraphs if p.text.strip()],
    'table_count': len(doc.tables),
    'table_data': []
}

# Extract table data
for table_idx, table in enumerate(doc.tables):
    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    data['table_data'].append(rows_data)

# Save as JSON
with open('docx_structured.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"Extracted {len(data['paragraphs'])} paragraphs and {len(data['table_data'])} tables")
